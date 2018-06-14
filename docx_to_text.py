# encoding=utf8
import os
import docx
import pandas as pd
from docx import Document

# in order to read from a folder anywhere in the dirve and write in the same folder
path = '/home/mohit/Desktop'
folder_name = 'word_docs'
folder_path = os.path.join(path,folder_name)
word_files = os.listdir(folder_path)
# print(word_files)

def getText(filename):
    file_path = os.path.join(folder_path,filename[0:-5])
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    data = '\n'.join(fullText)
    with open(file_path+'.txt','wb') as fi:
        fi.write(data.encode('utf-8'))

def getTables(filename):
    file_path = os.path.join(folder_path, filename[0:-5])
    doc = docx.Document(filename)
    all_tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # print cell.text
                row_data.append(cell.text)
            table_data.append(row_data)
        all_tables.append(table_data)

    # print(all_tables)
    for i,tab in enumerate(all_tables):
        df = pd.DataFrame(tab)
        df.to_csv(file_path+'-table-'+str(i)+'.csv',index=False)

    # print(df)

# filename = 'Meeting notes.docx'
for filename in word_files:
    getText(filename)
    getTables(filename)